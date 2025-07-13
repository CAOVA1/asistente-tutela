
import streamlit as st
from docx import Document

# Estado inicial de la conversación
if 'paso' not in st.session_state:
    st.session_state.paso = 0
    st.session_state.respuestas = {}

# Preguntas y llaves
preguntas = [
    ("¿Cuál es su nombre completo?", "nombre"),
    ("¿Cuál es su número de cédula?", "cedula"),
    ("¿Cuál es su dirección?", "direccion"),
    ("¿Cuál es su número de teléfono?", "telefono"),
    ("¿Cuál es su EPS?", "eps"),
    ("¿Qué diagnóstico tiene o qué enfermedad le afecta?", "diagnostico"),
    ("¿Qué medicamentos le formularon?", "medicamentos"),
    ("¿Qué procedimiento, medicamento o tratamiento le ha sido negado por la EPS?", "negacion"),
    ("¿Qué desea que el juez ordene a la EPS?", "pretensiones"),
]

def generar_documento(respuestas):
    doc = Document()
    doc.add_heading("ACCIÓN DE TUTELA", 0)
    doc.add_paragraph("Señor(a)\nJuez de la República\n(EPS correspondiente)\nCiudad")

    doc.add_heading("1. IDENTIFICACIÓN DEL ACCIONANTE", level=1)
    doc.add_paragraph(f"Nombre: {respuestas['nombre']}\nCédula: {respuestas['cedula']}\nDirección: {respuestas['direccion']}\nTeléfono: {respuestas['telefono']}")

    doc.add_heading("2. HECHOS", level=1)
    doc.add_paragraph(
        f"El accionante, afiliado a la EPS {respuestas['eps']}, ha sido diagnosticado con {respuestas['diagnostico']}. "
        f"Como parte de su tratamiento, le fueron formulados los siguientes medicamentos: {respuestas['medicamentos']}. "
        f"Sin embargo, la EPS ha negado el acceso a: {respuestas['negacion']}."
    )

    doc.add_heading("3. FUNDAMENTOS DE DERECHO", level=1)
    doc.add_paragraph(
        "Esta acción se fundamenta en los artículos 11 (derecho a la vida) y 49 (derecho a la salud) de la Constitución Política, "
        "así como en la Ley 1751 de 2015 y jurisprudencia constitucional relevante, como la Sentencia T-760 de 2008, que protegen el derecho fundamental a la salud."
    )

    doc.add_heading("4. PRETENSIONES", level=1)
    doc.add_paragraph(respuestas['pretensiones'])

    doc.add_heading("5. PRUEBAS", level=1)
    doc.add_paragraph("Historia clínica, fórmulas médicas, constancia de negación del servicio por la EPS.")

    doc.add_heading("6. JURAMENTO", level=1)
    doc.add_paragraph("Bajo la gravedad del juramento, manifiesto que no he presentado otra acción de tutela con el mismo objeto y finalidad.")

    doc.add_paragraph(f"\nAtentamente,\n\n[Firma]\n{respuestas['nombre']}\nC.C. {respuestas['cedula']}")
    doc.save("tutela_chat.docx")

# UI conversacional
st.title("🧑‍⚖️ Asesor Legal Virtual – Generador de Acción de Tutela")

if st.session_state.paso < len(preguntas):
    pregunta, clave = preguntas[st.session_state.paso]
    respuesta = st.text_input(pregunta, key=clave)

    if respuesta:
        st.session_state.respuestas[clave] = respuesta
        st.session_state.paso += 1
        st.rerun()

else:
    st.success("✅ ¡Información completa! Generando documento...")
    generar_documento(st.session_state.respuestas)

    with open("tutela_chat.docx", "rb") as file:
        st.download_button("📥 Descargar tutela generada", file, file_name="tutela_chat.docx")

    if st.button("🔁 Volver a empezar"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

